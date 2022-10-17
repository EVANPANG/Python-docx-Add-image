from dataclasses import dataclass
import numpy as np

import os
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

address = sys.argv[1]
sample = int(sys.argv[2])
quantity = int(sys.argv[3])

RowCount = (sample * quantity) / 2
# local = os.path.abspath('test.py')[::-1]
# address = local.split("yp.tset")[1][::-1]
data = []
doc = Document()
doc.styles['Normal'].font.size = Pt(10.5)
doc.styles['Normal'].font.name = u'仿宋'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')


def addTable(row, col):
    table = doc.add_table(rows=row, cols=col)
    return table


def walkFile(file):
    i = sample * quantity
    row = int(i / 2)
    table = addTable(row, 2)
    count = 1
    x = 0
    y = 0
    z = 1
    for root, dirs, files in os.walk(file):
        for f in files:
            try:
                data.append(f)
                cell = table.cell(x, y)
                text = cell.paragraphs[0]
                pic = text.add_run('')
                pic.add_picture(file + f,
                                width=Cm(5.29), height=Cm(3.53))
                print(f)
                pic.add_break()
                text.add_run('检材' + str(count) + '外观图')
                if z < quantity:
                    print(z)
                    count = count
                    z = z + 1
                else:
                    count = count + 1
                    z = 1
                if y == 1:
                    y = 0
                    x = x + 1
                else:
                    y = 1
            except:
                print('fail')


def main():
    walkFile(address + '/')
    doc.save(address + '/image.docx')


main()
